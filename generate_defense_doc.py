from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    colors = {
        1: (0, 70, 127),    # dark blue
        2: (0, 112, 192),   # medium blue
        3: (31, 73, 125),   # navy
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=16 if level==1 else 13 if level==2 else 11,
             bold=True, color=colors.get(level, (0,0,0)))
    return p

def add_body(doc, text, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_qa(doc, q_num, question, answer):
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"Q{q_num}. {question}")
    set_font(run, bold=True, size=11, color=(0, 70, 127))

    # Answer
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.25)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(answer)
    set_font(run2, size=10.5, italic=False)
    return p, p2

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, bold=True, size=10, color=(255, 255, 255))
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "00467F")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "EEF4FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=9.5)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 90)
    set_font(run, size=8, color=(180, 180, 180))


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("🎓 EduPredict AI")
set_font(run, size=28, bold=True, color=(0, 70, 127))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Final Year Project Defense Preparation Guide")
set_font(run2, size=16, bold=True, color=(0, 112, 192))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Student Academic Performance Prediction System")
set_font(run3, size=13, italic=True, color=(80, 80, 80))

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("Stack: Vue 3 · FastAPI · Scikit-Learn · SQLite")
set_font(run4, size=11, color=(100, 100, 100))

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("Defense Date: 17 June 2026")
set_font(run5, size=11, bold=True, color=(0, 70, 127))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "TABLE OF CONTENTS", 1)
toc_items = [
    ("PART 1", "Detailed Understanding of Your Implementation", "3"),
    ("  1.1", "What Is Your System In One Sentence", "3"),
    ("  1.2", "System Architecture – The Big Picture", "3"),
    ("  1.3", "The Machine Learning Pipeline – Step by Step", "4"),
    ("  1.4", "Derived Feature Computation", "6"),
    ("  1.5", "Prediction Pipeline – Single Student", "7"),
    ("  1.6", "Prediction Pipeline – Batch", "7"),
    ("  1.7", "Explainability Engine", "8"),
    ("  1.8", "API Endpoints", "8"),
    ("  1.9", "Database Design", "9"),
    ("  1.10", "Frontend Architecture", "9"),
    ("PART 2", "How Your System Solves the Existing Problem", "10"),
    ("PART 3", "100 Possible Defense Questions & Model Answers", "11"),
    ("  Section A", "Project Overview & Motivation (Q1–Q10)", "11"),
    ("  Section B", "Machine Learning & Model Design (Q11–Q35)", "13"),
    ("  Section C", "System Design & Architecture (Q36–Q55)", "17"),
    ("  Section D", "Data & Dataset (Q56–Q65)", "21"),
    ("  Section E", "Validation, Testing & Quality (Q66–Q75)", "23"),
    ("  Section F", "Literature & Theory (Q76–Q88)", "24"),
    ("  Section G", "Technical Deep-Dives (Q89–Q100)", "26"),
    ("APPENDIX", "Quick Revision Cheat Sheet", "29"),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(f"{num}  ")
    set_font(run_num, bold=(not num.startswith("  ")), size=10.5, color=(0,70,127))
    run_title = p.add_run(title)
    set_font(run_title, size=10.5)
    run_pg = p.add_run(f"  {'.' * max(1, 70 - len(num) - len(title))}  {pg}")
    set_font(run_pg, size=10, color=(120,120,120))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 1 — IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 1: DETAILED UNDERSTANDING OF YOUR IMPLEMENTATION", 1)
add_divider(doc)

# 1.1
add_heading(doc, "1.1  What Is Your System, In One Sentence?", 2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    "EduPredict AI is a full-stack, web-based, machine-learning-powered decision-support platform "
    "that predicts whether a Computer Science student will finish a semester as Excellent, Average, "
    "or At-Risk — enabling academic advisors to intervene before exams happen."
)
set_font(run, italic=True, size=11)

# 1.2
add_heading(doc, "1.2  System Architecture — The Big Picture", 2)
add_body(doc, "The system is built on a decoupled client-server architecture with 5 distinct layers:")
add_table(doc,
    ["Layer", "Technology", "File(s)"],
    [
        ["Frontend (Presentation)", "Vue 3, Vite, Axios, ChartJS, Tailwind CSS", "client/src/"],
        ["API Gateway", "FastAPI + Uvicorn (Python)", "server/app.py"],
        ["Routing", "FastAPI APIRouter", "routes/predict.py, records.py, insights.py, stats.py"],
        ["ML Engine", "Scikit-Learn (Random Forest)", "services/ml_service.py, train.py"],
        ["Explainability", "Rule-based heuristics", "services/explanation_service.py"],
        ["Analytics", "Pandas + SQLite queries", "services/stats_service.py"],
        ["Data Layer", "SQLite via Repository Pattern", "server/database.py"],
        ["Validation", "Pydantic v2 models", "server/schemas.py"],
    ],
    col_widths=[1.6, 2.2, 2.8]
)

# 1.3
add_heading(doc, "1.3  The Machine Learning Pipeline — Step by Step", 2)
steps = [
    ("STEP 1: Data Loading",
     "Dataset: cs_dataset.xlsx (sheet: Model_Ready). Real Computer Science student records. "
     "Student_ID column is dropped — it is not a predictive feature."),
    ("STEP 2: Encoding Categorical Variables",
     "Gender → LabelEncoder (Female=0, Male=1). Socioeconomic_Status → OrdinalEncoder (Low=0, Medium=1, High=2). "
     "Semester → OrdinalEncoder (First=0, Second=1). Performance_Status (target) → LabelEncoder (At-Risk=0, Average=1, Excellent=2). "
     "OrdinalEncoder is used for SES because Low < Medium < High has a meaningful rank order."),
    ("STEP 3: The 17 Features",
     "7 are direct user inputs: Gender, Age, Socioeconomic_Status, Level, Semester, Study_Hours_Per_Week, Previous_CGPA. "
     "10 are automatically derived by the backend from raw scores (see Section 1.4)."),
    ("STEP 4: Train-Test Split",
     "80% training / 20% testing. stratify=y ensures class distribution is preserved. random_state=42 ensures reproducibility."),
    ("STEP 5: Feature Scaling",
     "StandardScaler: x_scaled = (x - μ) / σ. Fit ONLY on training data, applied to both train and test. "
     "Tree models don't need scaling, but SVM and Logistic Regression do. Scaling ensures fair comparison."),
    ("STEP 6: Four Algorithms Trained",
     "Random Forest (n_estimators=200), Gradient Boosting (n_estimators=200, lr=0.1), "
     "Logistic Regression (max_iter=1000), Support Vector Machine (SVC with probability=True). "
     "All use class_weight='balanced' where supported."),
    ("STEP 7: Model Selection",
     "The winner is selected by highest WEIGHTED F1-SCORE on the test set — not accuracy — because accuracy is "
     "misleading with class imbalance."),
    ("STEP 8: Feature Importance",
     "Tree models: feature_importances_ attribute. Logistic Regression: absolute mean of coef_. "
     "Stored in insights.json."),
    ("STEP 9: Save 7 Artifacts",
     "best_model.pkl, scaler.pkl, le_status.pkl, le_gender.pkl, oe_ses.pkl, oe_semester.pkl, feature_cols.json, insights.json — all to server/models/."),
]
for title, body in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run_t = p.add_run(f"► {title}:  ")
    set_font(run_t, bold=True, size=11, color=(0,70,127))
    run_b = p.add_run(body)
    set_font(run_b, size=10.5)

add_heading(doc, "The 17 Features Used by the Model", 3)
add_table(doc,
    ["#", "Feature", "Type", "Source"],
    [
        ["1",  "Gender",                  "Demographic",  "Direct user input"],
        ["2",  "Age",                     "Demographic",  "Direct user input"],
        ["3",  "Socioeconomic_Status",    "Socio-economic","Direct user input"],
        ["4",  "Level",                   "Academic",     "Direct user input"],
        ["5",  "Semester",                "Academic",     "Direct user input"],
        ["6",  "Study_Hours_Per_Week",    "Behavioral",   "Direct user input"],
        ["7",  "Previous_CGPA",           "Academic",     "Direct user input"],
        ["8",  "Previous_Course_Average", "Derived",      "Avg of previous_scores"],
        ["9",  "Lowest_Previous_Score",   "Derived",      "Min of previous_scores"],
        ["10", "Weak_Previous_Count",     "Derived",      "Count of scores < 45"],
        ["11", "Course_CA_Average",       "Derived",      "Avg of ALL CA scores"],
        ["12", "Lowest_CA_Score",         "Derived",      "Min of all CA scores"],
        ["13", "Weak_CA_Count",           "Derived",      "Count of CA scores < 15"],
        ["14", "Core_CA_Average",         "Derived",      "Avg of core course CAs"],
        ["15", "Elective_CA_Average",     "Derived",      "Avg of elective CAs"],
        ["16", "University_CA_Average",   "Derived",      "Avg of university-wide CAs"],
        ["17", "Total_Units",             "Derived",      "Total courses registered"],
    ],
    col_widths=[0.3, 2.2, 1.4, 2.6]
)

# 1.4
add_heading(doc, "1.4  Derived Feature Computation — The Smart Part", 2)
add_body(doc,
    "When a user submits student scores, compute_derived_features() in ml_service.py "
    "automatically computes 10 features the user never manually calculates:")
derived = [
    ("Previous_Course_Average", "Mean of all previous course scores"),
    ("Lowest_Previous_Score",   "The minimum (worst) previous course score"),
    ("Weak_Previous_Count",     "Count of previous scores below 45 (the pass mark)"),
    ("Course_CA_Average",       "Mean of ALL current CA scores (core + elective + university + NUC)"),
    ("Lowest_CA_Score",         "The minimum current CA score"),
    ("Weak_CA_Count",           "Count of CA scores below 15 (danger threshold)"),
    ("Core_CA_Average",         "Mean of core Computer Science course CAs"),
    ("Elective_CA_Average",     "Mean of elective course CAs"),
    ("University_CA_Average",   "Mean of university-wide course CAs"),
    ("Total_Units",             "Total number of courses (len of all CA score lists)"),
]
for feat, desc in derived:
    add_bullet(doc, f"{feat}: {desc}")

add_body(doc,
    "SPECIAL CASE — 100 Level First Semester (\"100 Alpha\"): Fresh students have no prior university "
    "records. The system uses Entry_Academic_Score (e.g., JAMB score) as the previous-score proxy. "
    "previous_scores must be empty; Entry_Academic_Score is required.",
    bold=False, italic=True)

# 1.5
add_heading(doc, "1.5  Prediction Pipeline — Single Student", 2)
steps_single = [
    "User submits the form on the Vue frontend",
    "Pydantic (StudentInput) validates all inputs — types, ranges, 100-Alpha rules",
    "compute_derived_features() calculates the 10 derived values",
    "MLService._encode_row() converts categorical strings to numbers",
    "StandardScaler.transform() scales the 17-feature vector",
    "best_model.predict() → encoded numeric class label",
    "le_status.inverse_transform() → 'Excellent' / 'Average' / 'At-Risk'",
    "best_model.predict_proba() → probability scores for each class (confidence %)",
    "ExplanationService.build_explanation() → top 5 positive/negative reasons",
    "PredictionRepository.insert_prediction() → saved to SQLite database",
    "JSON response returned to frontend → Vue renders result card",
]
for i, s in enumerate(steps_single, 1):
    add_bullet(doc, f"Step {i}: {s}")

# 1.6
add_heading(doc, "1.6  Prediction Pipeline — Batch Upload", 2)
steps_batch = [
    "Advisor uploads .csv or .xlsx file via drag-and-drop on the frontend",
    "Frontend sends multipart POST to /predict/batch",
    "Backend reads file via Pandas (pd.read_csv or pd.read_excel)",
    "Normalize: 'Alpha'→'First', 'Omega'→'Second', 'Middle'→'Medium'",
    "Validate all 17 required columns exist",
    "Per-row validation: Gender, Level, Semester, score bounds (up to 20 errors returned)",
    "MLService.predict_batch() encodes + scales entire DataFrame at once",
    "model.predict() runs on all rows simultaneously (vectorized)",
    "Confidence scores computed per row via predict_proba()",
    "Results saved to SQLite (append or replace depending on 'replace' toggle)",
    "Return CSV download to browser via Content-Disposition header",
]
for i, s in enumerate(steps_batch, 1):
    add_bullet(doc, f"Step {i}: {s}")

# 1.7
add_heading(doc, "1.7  Explainability Engine (ExplanationService)", 2)
add_body(doc,
    "The ExplanationService is a rule-based heuristic system — NOT machine learning. "
    "It generates human-readable reasons by checking 8 threshold conditions:")
add_table(doc,
    ["Feature", "Positive Trigger", "Negative Trigger"],
    [
        ["Previous CGPA",         ">= 4.0",          "< 2.0"],
        ["Study Hours / Week",    ">= 20 hrs",        "< 8 hrs"],
        ["CA Average",            ">= 22 / 30",       "< 14 / 30"],
        ["Weak CA Count",         "—",                ">= 3 courses below 15"],
        ["Weak Previous Count",   "—",                ">= 2 past courses below 45"],
        ["Previous Course Avg",   ">= 70 / 100",      "< 45 / 100"],
        ["Core CA Average",       ">= 22 / 30",       "< 12 / 30"],
        ["Socioeconomic Status",  "High SES",         "Low SES (neutral warning)"],
    ],
    col_widths=[2.0, 2.0, 2.6]
)
add_body(doc, "The engine returns the top 5 most relevant explanations from all triggered checks.")

# 1.8
add_heading(doc, "1.8  API Endpoints", 2)
add_table(doc,
    ["Method", "Endpoint", "Purpose"],
    [
        ["GET",  "/health",        "Health check — confirms API is running"],
        ["POST", "/predict",       "Single student prediction"],
        ["POST", "/predict/batch", "Batch CSV/Excel file prediction"],
        ["GET",  "/records",       "Fetch saved prediction history from SQLite"],
        ["GET",  "/insights",      "Model metrics + feature importance rankings"],
        ["GET",  "/stats",         "Dashboard KPIs + chart data"],
    ],
    col_widths=[0.9, 1.8, 3.9]
)

# 1.9
add_heading(doc, "1.9  Database Design", 2)
add_body(doc,
    "SQLite database (students.db), single table: predictions. "
    "Uses a schema_version table to handle migrations without breaking the application. "
    "The PredictionRepository class implements the Repository Pattern — separating "
    "all SQL operations from business logic for clean, testable code.")
add_table(doc,
    ["Column", "Data Type", "Description"],
    [
        ["id",                    "INTEGER PK",  "Auto-increment unique record ID"],
        ["timestamp",             "DATETIME",    "When prediction was made (auto)"],
        ["student_id",            "TEXT",        "Optional tracking ID (nullable)"],
        ["gender",                "TEXT",        "Male / Female"],
        ["age",                   "INTEGER",     "Student age (15–60)"],
        ["socioeconomic_status",  "TEXT",        "Low / Medium / High"],
        ["level",                 "INTEGER",     "100 / 200 / 300 / 400"],
        ["semester",              "TEXT",        "First / Second"],
        ["study_hours_per_week",  "REAL",        "0.0 to 80.0"],
        ["previous_cgpa",         "REAL",        "0.00 to 5.00 (Nigerian 5-pt scale)"],
        ["previous_course_average","REAL",       "Derived: avg of historical scores"],
        ["course_ca_average",     "REAL",        "Derived: avg of current CA scores"],
        ["total_units",           "INTEGER",     "Derived: total courses registered"],
        ["predicted_status",      "TEXT",        "Excellent / Average / At-Risk"],
        ["confidence",            "REAL",        "Prediction probability (0.0–1.0)"],
    ],
    col_widths=[1.9, 1.3, 3.4]
)

# 1.10
add_heading(doc, "1.10  Frontend Architecture", 2)
add_table(doc,
    ["Component", "Role"],
    [
        ["Vue 3 (Composition API)", "Reactive UI components and state management"],
        ["Vue Router",             "SPA navigation without page reloads"],
        ["Axios",                  "HTTP client for communicating with FastAPI"],
        ["ChartJS",                "Dashboard visualizations (bar, pie, line charts)"],
        ["Tailwind CSS",           "Utility-first styling with glassmorphism effects"],
        ["Vite",                   "Build tool / dev server with instant HMR"],
    ],
    col_widths=[2.2, 4.4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 — HOW IT SOLVES THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 2: HOW YOUR SYSTEM SOLVES THE EXISTING PROBLEM", 1)
add_divider(doc)

add_heading(doc, "2.1  The Existing Problem — 5 Core Deficiencies", 2)
add_table(doc,
    ["Deficiency", "Description"],
    [
        ["Retrospective Nature",
         "Students are flagged as failing AFTER exams — when it is too late to help them"],
        ["Siloed & Inaccessible Data",
         "Each lecturer keeps their own spreadsheet; advisors cannot see cross-course performance"],
        ["High Manual Cognitive Load",
         "Advisors manually compute risk profiles across spreadsheets for large cohorts"],
        ["No Multi-Dimensional Correlation",
         "Spreadsheets cannot model non-linear patterns across 17+ factors simultaneously"],
        ["No Explainability",
         "When a student is flagged, there is no structured diagnosis of WHY they are struggling"],
    ],
    col_widths=[2.2, 4.4]
)

add_heading(doc, "2.2  How EduPredict AI Solves Each Problem", 2)
add_table(doc,
    ["Problem", "EduPredict AI Solution"],
    [
        ["Retrospective → Proactive",
         "Predicts outcome DURING the semester using current CA scores and study habits — before final exams occur"],
        ["Siloed → Centralized",
         "All data entered through one unified form or batch file; stored in one SQLite database"],
        ["Manual Load → Automated",
         "Auto-computes 10 derived features from raw scores; no manual GPA math required by the advisor"],
        ["No Correlation → ML Model",
         "17-feature Random Forest discovers non-linear patterns across demographics, behavior, and academic history"],
        ["No Explainability → Heuristic Engine",
         "ExplanationService labels each prediction with specific positive/negative academic drivers"],
    ],
    col_widths=[2.2, 4.4]
)

add_heading(doc, "2.3  The Core Argument — Know This By Heart", 2)
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.3)
p.paragraph_format.space_before = Pt(4)
run = p.add_run(
    "\"The fundamental shift EduPredict AI makes is from REACTIVE REMEDIATION "
    "(helping students after they have already failed) to PROACTIVE INTERVENTION "
    "(identifying at-risk students while there is still time to help them). "
    "A student with a CA average of 11/30 and 5 study hours per week can be flagged "
    "in Week 8 — giving the advisor time to schedule counseling, assign a peer tutor, "
    "or waive a coursework submission — before the final exam in Week 16.\""
)
set_font(run, italic=True, bold=False, size=11, color=(0, 70, 127))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 3 — 100 Q&A
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 3: 100 POSSIBLE DEFENSE QUESTIONS & MODEL ANSWERS", 1)
add_divider(doc)

sections = [
    ("SECTION A: PROJECT OVERVIEW & MOTIVATION  (Q1–Q10)", [
        ("What is EduPredict AI and what problem does it solve?",
         "EduPredict AI is a web-based machine learning decision-support system that predicts whether a Computer Science student will be Excellent, Average, or At-Risk at the end of a semester. It solves the problem of retrospective academic monitoring in Nigerian universities — where students are only identified as failing AFTER the semester ends, leaving no time for meaningful intervention."),
        ("Why did you choose this topic for your final year project?",
         "Academic failure rates in Nigerian tertiary institutions are consistently high. Traditional monitoring systems only capture failure after the fact. I saw an opportunity to apply machine learning to a real, impactful educational problem — building a system that gives advisors actionable intelligence while there is still time to help a student."),
        ("Who are the intended users of this system?",
         "The primary users are academic advisors and department administrators responsible for monitoring student welfare. The system is designed for non-technical users — educators don't need to understand machine learning; they just enter student data and receive a clear prediction with explanations."),
        ("What is the target output of the system?",
         "The system outputs one of three classifications: Excellent, Average, or At-Risk. It also provides a confidence score (probability %) and a list of the top explanatory factors that drove the prediction — both positive and negative."),
        ("How is this different from a student's CGPA?",
         "CGPA is a lagging indicator — it tells you how a student has performed in the past. EduPredict AI is a leading indicator — it uses current-semester data (CA scores, study hours) combined with historical data to predict future semester performance before the final exam happens."),
        ("What data does the system need to make a prediction?",
         "Seven direct inputs: Gender, Age, Socioeconomic Status, Level, Semester, Study Hours per Week, and Previous CGPA. Plus raw course scores (previous grades and current CA scores). The system then automatically derives 10 additional features from those raw scores."),
        ("Is the system specific to Computer Science students?",
         "Yes. The model was trained on a CS student dataset with CS-specific course structures. However, the architecture is generalizable and could be adapted for any department by retraining with appropriate data."),
        ("What are the three prediction outcomes and what do they mean?",
         "Excellent: the student is predicted to perform strongly and is not at risk. Average: performing adequately but with room for improvement. At-Risk: shows indicators of academic difficulty and needs immediate intervention."),
        ("How was the project scoped for a final year project?",
         "Scope was defined to focus on CS students at the 100–400 level using real student datasets. Out of scope: real-time LMS integration, attendance tracking, and multi-departmental models — documented as future work."),
        ("What real-world impact could this system have if deployed?",
         "Advisors could process an entire semester cohort's CA data mid-semester and immediately identify every At-Risk student. This enables proactive interventions — potentially reducing failure rates and student dropout significantly."),
    ]),
    ("SECTION B: MACHINE LEARNING & MODEL DESIGN  (Q11–Q35)", [
        ("What machine learning algorithms did you evaluate and why?",
         "Four algorithms: Random Forest, Gradient Boosting, Logistic Regression, and SVM. Chosen to cover a range: a linear baseline (LR), powerful ensembles (RF, GB), and a kernel-based method (SVM) — ensuring fair competitive evaluation."),
        ("Which algorithm won and why do you think it performed best?",
         "The best model is selected dynamically based on highest weighted F1-score on the test set. Random Forest and Gradient Boosting typically win on structured tabular data because they are ensemble methods that combine many weak learners and handle non-linear feature interactions well."),
        ("Why did you use F1-score as the selection metric instead of accuracy?",
         "Accuracy is misleading when class distributions are imbalanced. If 70% of students are Average, a model that always predicts 'Average' gets 70% accuracy but completely fails At-Risk detection. Weighted F1-score balances Precision and Recall across all classes."),
        ("Explain Precision, Recall, and F1-score in your context.",
         "Precision: of all students predicted At-Risk, what percentage actually are? (avoids false alarms). Recall: of all truly At-Risk students, what percentage did we catch? (avoids missed cases). F1: harmonic mean of both. In education, Recall for At-Risk is most critical — missing a student in danger is more costly than a false alarm."),
        ("What is a Random Forest classifier?",
         "An ensemble of decision trees. Each tree is trained on a random subset of data (bootstrap sampling) and uses a random subset of features at each split. The final prediction is the majority vote across all trees. This randomness reduces overfitting and variance."),
        ("How many trees does your Random Forest use?",
         "n_estimators=200 — 200 decision trees are built and their predictions are aggregated by majority vote."),
        ("What is class_weight='balanced' and why did you use it?",
         "It automatically adjusts sample weights inversely proportional to class frequencies. At-Risk students (if fewer in the dataset) are given higher weight during training, preventing the model from ignoring the minority class."),
        ("What does random_state=42 mean?",
         "It sets a seed for the random number generator, ensuring the training process produces the same results every run. This is essential for reproducibility — a key requirement in scientific research."),
        ("How did you split the data for training and testing?",
         "80% training, 20% testing using train_test_split with stratify=y. Stratification ensures the class distribution (ratio of Excellent:Average:At-Risk) is the same in both splits."),
        ("Why is stratification important?",
         "Without it, random splitting could place all At-Risk students in training and none in test, giving falsely optimistic results. Stratification guarantees representative sampling of all classes in both sets."),
        ("What is StandardScaler and why did you use it?",
         "StandardScaler standardizes features: x_scaled = (x - mean) / std. Applied because SVM and Logistic Regression are sensitive to feature scale. A feature in hundreds (previous scores 0-100) would otherwise dominate features in small numbers (study hours 0-80)."),
        ("Why do you fit the scaler only on training data?",
         "Fitting on test data causes data leakage — the model would peek at test statistics during training, leading to overly optimistic and unreliable performance estimates."),
        ("How does the system make a prediction at inference time?",
         "Input is encoded (categorical to numeric), derived features are computed, the 17-feature vector is constructed, the saved scaler.pkl transforms it, and best_model.pkl predicts the class and probability. le_status reverses the numeric prediction to a readable string."),
        ("What is predict_proba and how do you use it?",
         "predict_proba() returns the probability of each class. For example: Excellent: 72.3%, Average: 20.1%, At-Risk: 7.6%. The confidence score shown to the user is the probability of the predicted class."),
        ("How are feature importances computed?",
         "For tree-based models, feature_importances_ measures total reduction in Gini impurity contributed by each feature across all trees. For Logistic Regression, the absolute value of coefficients is used as a proxy."),
        ("What are 'weak previous counts' and 'weak CA counts'?",
         "Weak Previous Count: number of historical course scores below 45 (standard pass mark) — indicates a pattern of difficulty. Weak CA Count: number of current-semester CAs below 15 (out of 30) — an early warning signal that the student is already struggling."),
        ("Why is Previous_CGPA included when it already summarizes past performance?",
         "CGPA alone doesn't capture which specific areas a student is weak in. Combined with Weak_Previous_Count, Lowest_Previous_Score, and Previous_Course_Average, the model can distinguish between students with similar CGPAs but very different risk profiles."),
        ("What is the '100 Level First Semester' special case?",
         "Fresh 100-level students have no prior university records. The system requires an Entry_Academic_Score (e.g., JAMB/Post-UTME score) as substitute for previous_scores. This score computes Previous_Course_Average, Lowest_Previous_Score, and Weak_Previous_Count."),
        ("How does Gradient Boosting differ from Random Forest?",
         "Random Forest trains all trees independently and in parallel (bagging). Gradient Boosting trains trees sequentially — each new tree corrects the errors of the previous one. GB often achieves higher accuracy but is more prone to overfitting and slower to train."),
        ("What is overfitting and how did you guard against it?",
         "Overfitting: the model memorizes training data and fails on new data. Guards: train-test split for honest evaluation, class_weight='balanced', Random Forest's inherent bagging randomness, and using a held-out test set to select the best model."),
        ("Could you add more algorithms? Why did you stop at four?",
         "Yes — other options include KNN, Naive Bayes, or Neural Networks. The four chosen provide good balance of interpretability vs complexity. For this dataset size, deep neural networks would be overkill and harder to interpret for a defense audience."),
        ("What metrics are reported for each model?",
         "Accuracy, Precision (weighted), Recall (weighted), and F1-Score (weighted) — all computed on the 20% held-out test set. Results are stored in insights.json."),
        ("How is the model saved and loaded?",
         "joblib.dump() serializes the model to a .pkl binary file. joblib.load() deserializes it. Joblib is preferred over pickle for large NumPy arrays (like tree structures) — faster and more memory efficient."),
        ("What is the MLService Singleton pattern and why did you use it?",
         "MLService uses Python's __new__ method to ensure only one instance exists. This means .pkl files are loaded from disk ONCE when the server starts, not on every API request. Loading 7 model files on every prediction would add hundreds of milliseconds of latency."),
        ("What is Joblib and why use it instead of pickle?",
         "Joblib is optimized for serializing large scientific objects like NumPy arrays and Scikit-Learn models. It uses memory-mapped files for efficiency and is the Scikit-Learn recommended serialization method."),
    ]),
    ("SECTION C: SYSTEM DESIGN & ARCHITECTURE  (Q36–Q55)", [
        ("Explain your system architecture.",
         "Three-tier architecture: (1) Presentation — Vue 3 SPA frontend. (2) Application — FastAPI backend with four routers. (3) Data — SQLite database. The frontend communicates with the backend exclusively via HTTP JSON requests, making them fully decoupled."),
        ("Why did you choose FastAPI for the backend?",
         "FastAPI offers: automatic OpenAPI/Swagger documentation, native async support, Pydantic integration for automatic input validation, and is one of the fastest Python frameworks (comparable to NodeJS). It significantly reduces boilerplate compared to Flask or Django."),
        ("Why Vue 3 instead of React or Angular?",
         "Vue 3 with the Composition API offers reactive UI development with a gentle learning curve. Its Single File Components (.vue) keep HTML, JavaScript, and CSS organized. For a data-centric dashboard app, Vue's reactivity system is highly appropriate."),
        ("What is a Single Page Application (SPA)?",
         "An SPA loads one HTML page and dynamically updates content via JavaScript (Vue Router) without full page reloads. Navigation between Dashboard, Predictor, and Records feels instant — providing an app-like experience."),
        ("What is CORS and why did you configure it?",
         "CORS (Cross-Origin Resource Sharing) is a browser security mechanism blocking requests from a different origin. The frontend runs on localhost:5173 and backend on localhost:8000. Without CORS middleware, the browser blocks all API calls. allow_origins=['*'] permits all origins during development."),
        ("What does Pydantic do in your system?",
         "Pydantic's BaseModel defines StudentInput with field types, constraints, and validators. FastAPI automatically rejects out-of-range values, wrong types, or missing fields with a 422 response. This replaces hundreds of lines of manual validation code."),
        ("Explain the Repository Pattern used in your database layer.",
         "PredictionRepository in database.py encapsulates all SQL operations. Routes never write raw SQL — they call PredictionRepository.insert_prediction() or get_recent_records(). This separates data-access concerns from business logic, making code testable and maintainable."),
        ("Why SQLite instead of PostgreSQL or MySQL?",
         "For a final-year project, SQLite is ideal: zero server setup, the entire database is a single file (students.db), built into Python's standard library, and handles thousands of records. PostgreSQL would be the next step for production deployment."),
        ("Explain the database schema versioning system.",
         "The schema_version table tracks which schema is active. On startup, init_db() checks if version 2 is recorded. If not, it drops and recreates the predictions table with the new schema — preventing crashes when column structure changes between development iterations."),
        ("What are the four routers and what does each handle?",
         "predict.py: /predict (single) and /predict/batch. records.py: /records (retrieve logs). insights.py: /insights (model metrics, feature importances). stats.py: /stats (dashboard KPIs and chart data)."),
        ("How does the batch prediction return a downloadable file?",
         "The result DataFrame is converted to CSV string via to_csv(). FastAPI returns a Response with media_type='text/csv' and a Content-Disposition: attachment header. The browser automatically triggers a file download."),
        ("How does the system handle file uploads?",
         "FastAPI's UploadFile and File(...) handle multipart form data. The system reads file bytes into memory via await file.read(), then Pandas parses it with pd.read_csv() or pd.read_excel() based on the file extension."),
        ("What validation happens on a batch upload file?",
         "Three levels: (1) File type check — only .csv and .xlsx accepted. (2) Column check — all 17 required columns must be present. (3) Row-level validation — each row checked for valid Gender, Level, Semester, and score bounds. Up to 20 row errors returned at once."),
        ("How does your system normalize alternate data values?",
         "Some institutions use local terms: 'Alpha' for First semester, 'Omega' for Second, 'Middle' instead of 'Medium'. The batch route normalizes these before validation: df['Semester'] = df['Semester'].replace({'Alpha': 'First', 'Omega': 'Second'})."),
        ("What is Uvicorn and why is it used?",
         "Uvicorn is an ASGI web server. FastAPI is async and requires an ASGI server to run. Uvicorn is lightweight — suitable for both development (with --reload) and production deployment."),
        ("What does init_db() do when the server starts?",
         "It connects to SQLite, checks the schema version, drops and recreates the predictions table if the schema has changed, and ensures the table exists. Called once during server startup via create_app()."),
        ("How are ML artifacts loaded at server startup?",
         "MLService.__new__() ensures that when the first MLService() instance is created, load_models() is called once. All 7 .pkl and .json files are loaded into memory. All subsequent requests use the same in-memory objects."),
        ("What happens if train.py hasn't been run before starting the server?",
         "MLService.load_models() catches the exception and prints a warning. The server starts but predictions fail with a 503 error: 'ML models not loaded. Run train.py first.'"),
        ("How does your config file promote maintainability?",
         "config.py centralizes all constants: FEATURE_COLS, MAX_CA_SCORE, VALID_GENDERS, DB_FILE, MODEL_DIR. If any threshold or path changes, it's updated in one place — not scattered across multiple files."),
        ("What is the insights.json file used for?",
         "It stores training-time analytics: model comparison results (accuracy, precision, recall, F1 for all 4 models), feature importance rankings, class distribution percentages, and dataset info. The frontend's Insights dashboard reads this via the /insights endpoint."),
    ]),
    ("SECTION D: DATA & DATASET  (Q56–Q65)", [
        ("What dataset was used to train the model?",
         "A real Computer Science student dataset stored in cs_dataset.xlsx, sheet 'Model_Ready'. It contains student records with features like Gender, Level, Semester, Study Hours, CGPA, and derived CA metrics, with target label Performance_Status."),
        ("What is the target variable in your dataset?",
         "Performance_Status — a three-class label: Excellent, Average, or At-Risk. It represents the student's end-of-semester academic performance category."),
        ("How many records are in your dataset?",
         "The system prints the exact count during training: '[INFO] Loaded X records from sheet Model_Ready'. The count is visible in the training output and stored in insights.json under dataset_info."),
        ("What data preprocessing steps were applied?",
         "(1) Dropped Student_ID. (2) Dropped rows with missing Performance_Status. (3) Label-encoded Gender and Performance_Status. (4) Ordinal-encoded Socioeconomic_Status and Semester. (5) Converted Level to numeric. (6) Removed NaN rows from features. (7) Applied StandardScaler."),
        ("How did you handle class imbalance in the dataset?",
         "By using class_weight='balanced' for all applicable models and selecting the best model using weighted F1-score. The stratify=y parameter in train_test_split also ensures class ratios are maintained in both splits."),
        ("What is LabelEncoder vs OrdinalEncoder?",
         "LabelEncoder converts string labels to integers (0, 1, 2...) in alphabetical order — used for Gender and the target. OrdinalEncoder maps categories following a SPECIFIED order — used for SES (Low < Medium < High) and Semester (First < Second)."),
        ("Why is Student_ID dropped from training features?",
         "Because a student ID is an arbitrary identifier with no predictive relationship to academic performance. Including it would cause data leakage and make the model memorize individual students rather than learn generalizable patterns."),
        ("What does stratify=y mean in train_test_split?",
         "It ensures the proportion of each target class is the same in both training and testing sets as in the full dataset. This prevents a skewed split where, for example, all At-Risk students end up only in the training set."),
        ("Could you use this system with data from another department?",
         "Yes, with retraining. Replacing cs_dataset.xlsx with another department's data and running train.py produces new .pkl files. The API and frontend require no changes — demonstrating the system's reusability and modularity."),
        ("What is EDA and did you perform it?",
         "EDA (Exploratory Data Analysis) involves analyzing dataset distributions, correlations, missing values, and outliers before modeling. Yes — an EDA.ipynb Jupyter Notebook is included in the server directory. It informed feature selection and explanation threshold choices."),
    ]),
    ("SECTION E: VALIDATION, TESTING & QUALITY  (Q66–Q75)", [
        ("How did you test your system?",
         "Two levels: (1) Unit tests — test_app.py and test_api.py using Pytest and HTTPX test client to verify API endpoint responses. (2) Manual testing — using Swagger UI at http://127.0.0.1:8000/docs and the Vue frontend with real student records."),
        ("What is Pytest and how is it used here?",
         "Pytest is a Python testing framework. Test files create a TestClient from FastAPI's test utilities to make HTTP requests against the API without needing a running server. Each function validates a specific endpoint's response structure and status code."),
        ("How does Pydantic validation protect the API?",
         "StudentInput has @field_validator methods and range constraints (ge=, le=) on every field. FastAPI automatically rejects any failing request with a 422 Unprocessable Entity response and detailed error message — the server never reaches ML logic for invalid inputs."),
        ("What happens when a user submits invalid data?",
         "Pydantic raises a ValidationError. FastAPI catches it and returns a 422 response with a JSON body listing every field that failed validation and why. Example: 'Age must be between 15 and 60'."),
        ("How is the 100-Alpha validation rule implemented?",
         "A @model_validator(mode='after') in schemas.py runs after all field validators. If Level=100 AND Semester='First': Entry_Academic_Score is required and previous_scores must be empty. Otherwise: previous_scores must have at least one entry."),
        ("What is the confidence score and how reliable is it?",
         "The confidence is the model's predict_proba() value for the predicted class — e.g., 85% means the ensemble of 200 trees agrees 85% on that class. Higher confidence = more agreement = more reliable prediction."),
        ("What is the performance requirement for the API?",
         "Non-functional requirement: single prediction requests must return within 100 milliseconds. Achievable because the model is already in memory (no file I/O on each request) and inference on a 17-feature vector is computationally trivial for a tree model."),
        ("Did you handle privacy in the system?",
         "Yes. Student_ID is optional and used only for tracking — not a prediction feature. The system does not require or store names, phone numbers, or other PII beyond what is necessary for auditing."),
        ("What happens if the batch file has invalid rows?",
         "The system collects up to 20 validation errors from all rows and returns them all at once in a 422 response. This prevents the user from fixing one error, re-uploading, discovering another — improving usability significantly."),
        ("How did you verify that derived features match training data expectations?",
         "By ensuring compute_derived_features() in ml_service.py uses the identical logic and thresholds as the dataset preparation script. feature_cols.json also enforces that features are always passed in the correct order to the model."),
    ]),
    ("SECTION F: LITERATURE & THEORY  (Q76–Q88)", [
        ("What existing systems or research did you review?",
         "Existing research includes works using decision trees, neural networks, and logistic regression on datasets like the UCI Student Performance dataset (Cortez & Silva, 2008). Key finding: ensemble methods (Random Forest, Gradient Boosting) consistently outperform single models on tabular educational data."),
        ("What is the difference between classification and regression?",
         "Classification predicts a discrete category — Excellent, Average, or At-Risk. Regression predicts a continuous number like exact CGPA. Classification is more appropriate because academic policy decisions (intervention, counseling) are made based on categorical risk levels."),
        ("Why is this called a 'Decision Support System' rather than 'Automated Decision System'?",
         "EduPredict AI provides predictions and explanations to SUPPORT human decision-making. The final decision always rests with the academic advisor. The system does not automatically act on any student without human review."),
        ("What is the difference between supervised and unsupervised learning?",
         "Supervised learning trains on labeled data (each record has a known Performance_Status). Unsupervised learning finds patterns in unlabeled data (e.g., clustering). EduPredict AI uses supervised learning — the training dataset contains ground-truth outcome labels."),
        ("What is ensemble learning and why is it better than a single model?",
         "Ensemble learning combines predictions from multiple models for a more robust result. Random Forest averages 200 trees' votes, reducing the impact of any single tree's errors. The ensemble is less sensitive to noise and outliers than any individual tree."),
        ("What is Gini impurity in the context of decision trees?",
         "Gini impurity measures how 'mixed' a node is — a node with all samples from one class has Gini=0 (pure). Random Forest splits nodes to minimize Gini impurity. Feature importance measures each feature's contribution to total impurity reduction across all trees."),
        ("What is data leakage and how did you prevent it?",
         "Data leakage: information from outside training contaminates the model. Prevention: (1) StandardScaler fit only on training data. (2) Student_ID excluded from features. (3) Target column excluded from the feature matrix X."),
        ("Why is this a multi-class problem rather than binary?",
         "Reducing to binary (At-Risk vs Not At-Risk) would discard useful information. Knowing whether a student will be Average vs Excellent is valuable for prioritizing limited intervention resources — advisors should focus most effort on At-Risk, some on borderline Average, and less on Excellent."),
        ("What ethical considerations exist with an AI academic prediction system?",
         "Key concerns: (1) Algorithmic bias — if training data reflects historical biases (e.g., against low-SES students), the model may perpetuate them. (2) Self-fulfilling prophecy — labeling a student 'At-Risk' might affect how educators treat them. (3) Privacy of student data. Addressed through: human-in-the-loop design, optional student ID, and transparent explanations."),
        ("What are the limitations of your system?",
         "(1) The model is only as good as the training data — biases propagate. (2) It doesn't capture real-time behavioral data like attendance. (3) It predicts semester-end performance from mid-semester data. (4) SQLite limits concurrent users in production."),
        ("What future improvements could be made?",
         "(1) LMS integration (Moodle, Canvas) for automatic data ingestion. (2) Attendance tracking integration. (3) Email/SMS alerts for flagged students. (4) Semester-by-semester model retraining. (5) SHAP values for deeper explainability. (6) PostgreSQL for production scalability."),
        ("What is SHAP and why would it improve your system?",
         "SHAP (SHapley Additive exPlanations) assigns each feature a mathematically grounded contribution value for a specific prediction using game theory. Unlike the current heuristic threshold system, SHAP shows exactly how much each feature pushed the prediction in a given direction."),
        ("How does socioeconomic status affect academic predictions?",
         "Low-SES students may lack textbooks, stable internet, study space, or have part-time jobs. The model includes it as a feature, and the Explanation Engine flags it as a contextual factor — labeled 'neutral' rather than 'negative' to avoid stigmatization."),
    ]),
    ("SECTION G: TECHNICAL DEEP-DIVES  (Q89–Q100)", [
        ("Walk me through exactly what happens when I click 'Predict' on the frontend.",
         "(1) Vue validates form fields. (2) Axios sends POST /predict with JSON. (3) Pydantic validates all fields. (4) ml_service.predict_single() is called. (5) compute_derived_features() calculates 10 derived values. (6) The 17-feature row is encoded and scaled. (7) best_model.predict() returns the class index. (8) le_status.inverse_transform() converts to readable string. (9) predict_proba() gives confidence %. (10) ExplanationService generates reasons. (11) PredictionRepository saves to SQLite. (12) JSON response sent. (13) Vue renders the result card."),
        ("Why does your system save encoders as separate .pkl files?",
         "Because the encoders learned their mapping during training and must apply the EXACT same mapping at inference time. Saving them separately ensures perfect consistency — even if the server restarts or the script is rerun months later."),
        ("What would happen if you applied a different scaler at prediction time?",
         "The prediction would be completely wrong. StandardScaler shifts and scales each feature by the TRAINING SET's mean and standard deviation. A different scaler would place input values in a completely different region of the feature space, producing nonsense predictions."),
        ("How does the get_db() context manager work?",
         "The @contextmanager decorator makes get_db() work with 'with' statements. It opens a SQLite connection, yields it for use, then automatically closes it — even if an exception occurs. This prevents connection leaks."),
        ("What is sqlite3.Row and why do you set it as row_factory?",
         "sqlite3.Row makes query results accessible by column name (row['predicted_status']) in addition to index. Setting conn.row_factory = sqlite3.Row allows dict(row) to convert database rows directly to Python dictionaries — cleaner and less error-prone."),
        ("How does Vue Router work in your SPA?",
         "Vue Router maps URL paths (/dashboard, /predict, /records) to Vue components. When a user clicks a navigation link, Vue Router swaps the rendered component without making a new server request, creating seamless, fast navigation."),
        ("Why do you normalize 'Alpha' to 'First' in the batch route?",
         "The OrdinalEncoder was trained on 'First' and 'Second'. If a CSV uses 'Alpha' or 'Omega', the encoder would fail with an unknown category error. Normalization ensures data consistency regardless of the upload source's local terminology."),
        ("What is the 'replace' parameter in the batch upload endpoint?",
         "When replace=True, PredictionRepository.replace_predictions() deletes ALL existing records then inserts the new batch. When replace=False, it appends via insert_predictions(). This gives advisors flexibility to replace history or grow it incrementally."),
        ("How does Vite differ from traditional webpack?",
         "Vite uses native ES Modules in the browser during development — no bundling step. This makes the dev server start nearly instantly and enables extremely fast Hot Module Replacement (HMR). Webpack bundles everything first, making it slower for development."),
        ("Why use Axios instead of the native browser fetch API?",
         "Axios provides automatic JSON serialization/deserialization, cleaner error handling (errors throw exceptions rather than requiring manual status checks), request/response interceptors for global error handling, and easier multipart file upload syntax."),
        ("How would you scale this system for 10,000 concurrent users?",
         "(1) Replace SQLite with PostgreSQL. (2) Deploy FastAPI behind a load balancer with multiple Uvicorn workers (Gunicorn). (3) Cache model artifacts in Redis. (4) Use a CDN for Vue static files. (5) Containerize with Docker and orchestrate with Kubernetes."),
        ("If you had to redo this project, what would you do differently?",
         "(1) Implement SHAP for deeper explainability instead of manual heuristics. (2) Collect time-series data across multiple weeks. (3) Use PostgreSQL from the start. (4) Implement proper authentication/authorization. (5) Build an automated model retraining pipeline that updates the model each semester."),
    ]),
]

q_num = 1
for sec_title, qas in sections:
    doc.add_page_break()
    add_heading(doc, sec_title, 2)
    add_divider(doc)
    for question, answer in qas:
        add_qa(doc, q_num, question, answer)
        q_num += 1

# ══════════════════════════════════════════════════════════════════════════════
#  CHEAT SHEET
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "APPENDIX: QUICK REVISION CHEAT SHEET", 1)
add_divider(doc)
add_body(doc, "Know these numbers and facts cold before your defense:", bold=True)
add_table(doc,
    ["Topic", "Key Fact"],
    [
        ["Total features",              "17 features (7 direct input + 10 auto-derived)"],
        ["Target classes",              "Excellent · Average · At-Risk"],
        ["Algorithms compared",         "4 (Random Forest, Gradient Boosting, Logistic Regression, SVM)"],
        ["Model selection metric",      "Weighted F1-Score (NOT accuracy)"],
        ["Best model (typical)",        "Random Forest — 200 trees (n_estimators=200)"],
        ["Train / Test split",          "80% training / 20% testing (stratified)"],
        ["Scaling method",              "StandardScaler — fit on TRAINING data only (z-score: x_scaled = (x-μ)/σ)"],
        ["Database",                    "SQLite (students.db) — Repository Pattern"],
        ["Gender encoding",             "LabelEncoder → Female=0, Male=1"],
        ["SES encoding",                "OrdinalEncoder → Low=0, Medium=1, High=2"],
        ["Semester encoding",           "OrdinalEncoder → First=0, Second=1"],
        ["Target encoding",             "LabelEncoder → At-Risk=0, Average=1, Excellent=2"],
        ["Explainability method",       "Rule-based heuristics (8 threshold checks) — NOT ML"],
        ["API framework",               "FastAPI + Pydantic v2 + Uvicorn"],
        ["Frontend",                    "Vue 3 (Composition API) + Vite + Axios + ChartJS"],
        ["100-Alpha special case",      "Entry_Academic_Score used as proxy (no university history)"],
        ["Pass threshold (prev score)", "< 45 = weak course (Weak_Previous_Count)"],
        ["CA danger threshold",         "< 15 / 30 = weak CA (Weak_CA_Count)"],
        ["CA positive threshold",       ">= 22 / 30 = positive explanation"],
        ["Confidence score =",          "predict_proba() value of the predicted class"],
        ["Singleton pattern used for",  "MLService — loads model files once at startup"],
        ["Number of saved .pkl files",  "7 files: best_model, scaler, le_status, le_gender, oe_ses, oe_semester + feature_cols.json"],
        ["CGPA scale",                  "Nigerian 5-point scale (0.00 – 5.00)"],
        ["Max CA score per course",     "30 marks"],
        ["Max study hours field",       "80 hours/week"],
    ],
    col_widths=[2.6, 4.0]
)

doc.add_paragraph()
p_tip = doc.add_paragraph()
p_tip.paragraph_format.left_indent = Inches(0.2)
run_tip = p_tip.add_run(
    "DEFENSE TIP 1: If you don't know an exact number, say: "
    "'The exact figure is in the insights.json output from training, but the model was selected as best on the test set.' "
    "Never guess a number. Confidence in what you know matters more than knowing everything.\n\n"
    "DEFENSE TIP 2: Always tie technical answers back to real-world impact — "
    "'...and this matters because it means an advisor can identify a struggling student in Week 8, "
    "with 8 weeks still remaining to intervene before the final exam.'\n\n"
    "Good luck tomorrow! You built something genuinely impactful — own it with confidence."
)
set_font(run_tip, italic=True, size=11, color=(0, 70, 127))

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\Chiso\Desktop\Among_Us\finalyear_assignment2\EduPredict_Defense_Preparation_Guide.docx"
doc.save(output_path)
print(f"[SUCCESS] Word document saved to: {output_path}")
